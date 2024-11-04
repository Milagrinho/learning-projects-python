from docx import Document
from docx.shared import Inches
import pyttsx3

def speak(text):
    pyttsx3.speak(text)

document = Document()

document.add_heading('Curriculum Vitae', 0)

# Profile picture
document.add_picture(
    'Foto Perfil.png', 
    width=Inches(2.0)
)

# Name, phone number and email details
name = input('What is your name? ')
speak('Hello, ' + name + '. Hope you are having a great day!')
phone_number = input('What is your phone number? ')
email = input('What is your email? ')

document.add_paragraph( name + ' | ' + phone_number + ' | ' + email)

# About me
document.add_heading('About Me')
document.add_paragraph(
    input('Tell me about yourself? ')
)

# Work experiences
document.add_heading('Work experiences')
p = document.add_paragraph()

company = input('Enter company name: ')
from_date = input('From Date: ')
to_date = input('To Date: ')

p.add_run(company + ' ').bold = True
p.add_run(from_date + '-' + to_date + '\n').italic = True

experience_details = input(
    'Describe your experience at ' + company + ': ')
p.add_run(experience_details)

# More experiences
while True:
    has_more_experiences = input(
        'Do you have more experiences? Y/N ')
    if has_more_experiences.upper() == 'Y':
        p = document.add_paragraph()

        company = input('Enter company name: ')
        from_date = input('From Date: ')
        to_date = input('To Date: ')

        p.add_run(company + ' ').bold = True
        p.add_run(from_date + '-' + to_date + '\n').italic = True

        experience_details = input(
            'Describe your experience at ' + company + ': ')
        p.add_run(experience_details)
    else:
        break

# Skills
document.add_heading('My skills')
skill = input('What is your skill? ')
s = document.add_paragraph(skill)
s.style = 'List Bullet'

while True:
    has_more_skill = input(
        'Do you have more skills? Y/N ')
    if has_more_skill.upper() == 'Y':
        skill = input('What is your skill? ')
        s = document.add_paragraph(skill)
        s.style = 'List Bullet'
    else:
        break

# Footer
section = document.sections[0]
footer = section.footer
f = footer.paragraphs[0]
f.text = 'CV generated using Amigocodes and Intuit QuickBooks course project'

document.save('cv.docx')